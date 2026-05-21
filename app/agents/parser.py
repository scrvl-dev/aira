"""
Document Parser
Routes uploaded files to the correct parser based on filename + content sniffing.
Returns raw text per document for the field extractor.
"""
import io
import re
from pathlib import Path
from typing import Optional

# These are imported lazily so missing deps don't crash the whole app
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


DOC_TYPE_PATTERNS = {
    "submission": [
        "submission", "sub_sheet", "sub sheet"
    ],
    "works": [
        "list_of_works", "list of works", "works", "low",
        "ih_-_list", "ih - list"
    ],
    "survey": [
        "condition_survey", "condition survey", "building_condition",
        "building condition", "survey_report", "survey report"
    ],
    "questionnaire": [
        "questionnaire", "property_questionnaire", "mtr_property",
        "property questionnaire"
    ],
    "valuation": [
        "valuation", "val_report", "bv_real", "valuation_report"
    ],
}


def detect_doc_type(filename: str, content_sample: str = "") -> str:
    """Detect document type from filename first, then content."""
    name_lower = filename.lower().replace("-", "_").replace(" ", "_")
    
    for doc_type, patterns in DOC_TYPE_PATTERNS.items():
        for pattern in patterns:
            if pattern.replace(" ", "_") in name_lower:
                return doc_type
    
    # Fallback: content-based detection
    content_lower = content_sample.lower()
    if "submission sheet" in content_lower or "lender name" in content_lower:
        return "submission"
    if "list of essential works" in content_lower or "works identified" in content_lower:
        return "works"
    if "building condition survey" in content_lower or "condition rating" in content_lower:
        return "survey"
    if "mtr property questionnaire" in content_lower or "household composition" in content_lower:
        return "questionnaire"
    if "market value" in content_lower and "rebuilding cost" in content_lower:
        return "valuation"
    
    return "unknown"


def parse_pdf(file_bytes: bytes) -> str:
    """Extract all text from PDF using pdfplumber."""
    if not HAS_PDFPLUMBER:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[PAGE {i+1}]\n{page_text}")
            
            # Also extract any tables
            tables = page.extract_tables()
            for table in tables:
                if table:
                    table_text = "\n".join(
                        " | ".join(str(cell or "").strip() for cell in row)
                        for row in table if any(cell for cell in row)
                    )
                    if table_text.strip():
                        text_parts.append(f"[TABLE PAGE {i+1}]\n{table_text}")
    
    return "\n\n".join(text_parts)


def parse_excel(file_bytes: bytes) -> str:
    """Extract all data from Excel as structured text."""
    if not HAS_OPENPYXL:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")
    
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    text_parts = []
    
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        sheet_rows = []
        
        for row in ws.iter_rows(values_only=True):
            # Skip empty rows
            if not any(cell is not None for cell in row):
                continue
            # Format row as pipe-separated
            row_text = " | ".join(
                str(cell).strip() if cell is not None else ""
                for cell in row
            )
            if row_text.replace("|", "").strip():
                sheet_rows.append(row_text)
        
        if sheet_rows:
            text_parts.append(f"[SHEET: {sheet_name}]\n" + "\n".join(sheet_rows))
    
    return "\n\n".join(text_parts)


def parse_docx(file_bytes: bytes) -> str:
    """Extract text from Word document."""
    if not HAS_DOCX:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = DocxDocument(io.BytesIO(file_bytes))
    text_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())
    
    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells
            )
            if row_text.replace("|", "").strip():
                text_parts.append(row_text)
    
    return "\n".join(text_parts)


def parse_file(filename: str, file_bytes: bytes) -> tuple[str, str]:
    """
    Parse a file and return (doc_type, extracted_text).
    Auto-detects document type from filename + content.
    """
    filename_lower = filename.lower()
    ext = Path(filename).suffix.lower()
    
    # Parse based on extension
    if ext == ".pdf":
        raw_text = parse_pdf(file_bytes)
    elif ext in (".xlsx", ".xls"):
        raw_text = parse_excel(file_bytes)
    elif ext in (".docx", ".doc"):
        raw_text = parse_docx(file_bytes)
    elif ext == ".csv":
        raw_text = file_bytes.decode("utf-8", errors="replace")
    else:
        # Try PDF first, then Excel
        try:
            raw_text = parse_pdf(file_bytes)
        except Exception:
            raw_text = file_bytes.decode("utf-8", errors="replace")
    
    # Detect document type
    doc_type = detect_doc_type(filename, raw_text[:500])
    
    return doc_type, raw_text


def parse_batch(files: dict[str, bytes]) -> dict[str, str]:
    """
    Parse a batch of files.
    Input: {filename: file_bytes}
    Output: {doc_type: extracted_text}
    """
    results = {}
    
    for filename, file_bytes in files.items():
        try:
            doc_type, text = parse_file(filename, file_bytes)
            if doc_type in results:
                # If we get a duplicate type, use longer/better one
                if len(text) > len(results[doc_type]):
                    results[doc_type] = text
            else:
                results[doc_type] = text
        except Exception as e:
            results[f"error_{filename}"] = f"PARSE ERROR: {str(e)}"
    
    return results
